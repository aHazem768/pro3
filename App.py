#------------------------------------------------------------------------------------------------------------------------------------
#  Импортируйте все библиотеки
import streamlit as st
import pandas as pd
from datetime import datetime
from PIL import Image
import matplotlib.pyplot as plt
import numpy as np
import plotly.express as px
from time import sleep
import re
import os
from docx import Document
from docx.shared import Inches
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#  Запустите программу и разместите ее изображение
st.title(" 📊 Project Productivity Professional ")
st.markdown("---")
st.write('')
st.sidebar.title('***📊Добро пожаловать📊***')
st.sidebar.markdown("---")
image = Image.open("142.jpeg")
st.sidebar.image(image, caption='', use_column_width=True)
st.sidebar.markdown("---")
st.sidebar.title("Информация о проекте")
# Установка максимального количества ячеек для отображения в Pandas Styler
pd.set_option("styler.render.max_elements", 559776)
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
# Основные характеристики
if st.sidebar.checkbox("Основные характеристики"):
    selected_dashboards = st.multiselect("Основные характеристики Проекта", ["Основные характеристики"])
    show_fields = st.checkbox("Характеристики продукта ")
    inputs = {}  # Пустой словарь для хранения введенных значений
    if show_fields:
        for dashboard_option in selected_dashboards:
            if dashboard_option == "Основные характеристики":
                inputs["Название Проекта"] = st.text_input("Название Проекта", value="")
                inputs["Адрис Проекта"] = st.text_input("Адрис Проекта", value="")
                inputs["Вид работ"] = st.text_input("Вид работ", value="")
                inputs["Генпроектировщик"] = st.text_input("Генпроектировщик", value="")
                inputs["Номер договора"] = st.text_input("Номер договора", value="")
                inputs["Закачик"] = st.text_input("Закачик", value="")
                inputs["Стадия проекта"] = st.text_input("Стадия Проекта", value="")
                inputs["Назначение Объекта"] = st.text_input("Назначение Объекта", value="")
                inputs["Номер ГПЗУ"] = st.text_input("Номер ГПЗУ", value="")
                inputs["Этажность"] = st.text_input("Этажность", value="")
                inputs["Количество Надземных Этажей"] = st.text_input("Количество Надземных Этажей", value="")
                inputs["Количество Подземных Этажей"] = st.text_input("Количество Подземных Этажей", value="")
                inputs["Количество секций"] = st.text_input("Количество секций", value="")
                inputs["Общая площадь здания m²"] = st.text_input("Общая площадь здания m² ", value="")
                inputs["Строительный Объем m³"] = st.text_input("Строительный Объем m³", value="")  
                inputs["Площадь участка"] = st.text_input("Площадь участка", value="")
                inputs["Полезная площадь"] = st.text_input("Полезная площадь", value="")
    if len(inputs) > 0:  # Просматривайте информационную таблицу только при наличии сохраненных данных
        st.write("***Основные характеристики Проекта:***")
        with st.expander("Показать таблицу"):
            df = pd.DataFrame(inputs.items(), columns=['Имя атрибута', 'Описание'])
            st.table(df)   
            #  Создайте кнопку загрузки таблицы в файл Excel
            if st.checkbox('скачать таблицу в формате Excel'):
                df_to_save = pd.DataFrame(inputs.items(), columns=['Имя атрибута', 'Описание'])
                excel_file = 'project_characteristics.xlsx'
                with pd.ExcelWriter(excel_file) as writer:
                    df_to_save.to_excel(writer, index=False, sheet_name='Основные характеристики Проекта')
                    st.markdown(f'<a href="{excel_file}" download="project_characteristics.xlsx">Нажмите здесь, чтобы скачать таблицу в формате Excel</a>', unsafe_allow_html=True)

#-------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------
#--------------------------------------------------------------------------------------------------------------------------------------------------

# Кнопки "Разработчики" و "График" для главного выключателя
st.sidebar.title("Расчет трудоемкости")
show_upload_button = st.sidebar.checkbox("Разработчики")
# Добавление кнопки для отображения Просмотр каждого разработчика
show_dashboard = st.sidebar.checkbox("Просмотр каждого разработчика")
show_chart_button = st.sidebar.checkbox("Показать график")
# Загрузка файла журнала
if show_upload_button:
    st.write('')
    uploaded_files = st.sidebar.file_uploader("Выберите файл журнала (Log file)", type=['log'])  # Кнопка загрузки файла Log
    if uploaded_files is not None:  
        # Чтение данных из файла Log
        log_content = uploaded_files.readlines()

        # Преобразуйте содержимое в строки
        lines = [line.decode().strip() for line in log_content]

        # Разделите строки и удалите пустые строки
        lines = [line.strip() for line in lines if line.strip()]

        # Разделите каждую строку на две части с помощью вертикальной черты
        data = [line.split('|')[:2] for line in lines if '|' in line]

        # Создание DataFrame
        df = pd.DataFrame(data, columns=['Дата и время', 'Разработчики'])
        show_upload_button = st.sidebar.checkbox("Общий протокол")
        if show_upload_button:
          st.write('Общий протокол')
          df
        # Применение шаблона к именам пользователей
        pattern = re.compile(r'^[А-Яа-я]+\s[А-Яа-я]+$|^\w+\s\w+$')  
        df = df[df['Разработчики'].apply(lambda x: bool(pattern.match(x.strip())))]  
        
        # Объединение столбцов 'Дата' и 'время' в новый столбец 'Общее время'
        df['Общее время'] = pd.to_datetime(df['Дата и время'], errors='coerce')
        
        # Группировка данных по столбцу 'Разработчики'
        grouped_data = df.groupby('Разработчики')
        
        # Создание пустого списка для хранения результатов (с и без времени отдыха)
        results = []
        cleaned_results = []

        # Вычисление общего времени для каждого разработчика (включая исключение времени отдыха < 30 минут)
        for name, group in grouped_data:
            # Вычисляем разницу во времени
            time_diffs = group['Общее время'].diff().fillna(pd.Timedelta(seconds=0))

            # Исключаем периоды отрицательного времени
            time_diffs = time_diffs[time_diffs >= pd.Timedelta(seconds=0)]

            # Вычисляем общее время после исключения отрицательного времени
            total_time_seconds = time_diffs.sum().total_seconds()

            # Преобразование времени в дни, часы, минуты и секунды
            days, remainder = divmod(total_time_seconds, 86400)
            hours, remainder = divmod(remainder, 3600)
            minutes, seconds = divmod(remainder, 60)

            # Форматирование времени
            formatted_time = f"{int(days)} days, {int(hours)} hours, {int(minutes)} minute, {int(seconds)} seconds"

            # Добавление результатов в список
            results.append({'Разработчики': name, 'Время': formatted_time})

            # Исключаем периоды времени отдыха более 30 минут
            time_diffs = time_diffs[time_diffs <= pd.Timedelta(minutes=30)]

            # Вычисляем общее время после исключения отрицательного времени и времени отдыха
            total_time_seconds = time_diffs.sum().total_seconds()

            # Преобразование времени в дни, часы, минуты и секунды
            days, remainder = divmod(total_time_seconds, 86400)
            hours, remainder = divmod(remainder, 3600)
            minutes, seconds = divmod(remainder, 60)

            # Форматирование времени
            formatted_time = f"{int(days)} days, {int(hours)} hours, {int(minutes)} minute, {int(seconds)} seconds"

            # Добавление результатов в список
            cleaned_results.append({'Разработчики': name, 'Время': formatted_time})
        
        # Создание DataFrame из результатов с и без времени отдыха
        results_df = pd.DataFrame(results)
        cleaned_results_df = pd.DataFrame(cleaned_results)

        # Вывод результатов в виде таблиц
        st.write('Трудоемкость разработчиков:')
        st.write(results_df)
        st.write('Трудоемкость разработчиков (без времени отдыха < 30 минут):')
        st.write(cleaned_results_df)
        
 
        # Добавление кнопки для отображения Просмотр каждого разработчика
        if show_dashboard:
            st.markdown("---")
            # Получение списка всех имен сотрудников
            developers = cleaned_results_df['Разработчики'].unique()

            # Создание выпадающего списка для выбора сотрудника
            selected_developers = st.multiselect("Выберите разработчиков", developers)

            # Фильтрация данных для выбранных сотрудников
            if selected_developers:
                filtered_data = cleaned_results_df[cleaned_results_df['Разработчики'].isin(selected_developers)]
                for developer in selected_developers:
                    st.write(f"Сводная таблица для {developer}:")
                    st.write(filtered_data[filtered_data['Разработчики'] == developer])
            
#-----------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------



if show_chart_button:
    # Просмотр параметров для типа диаграммы
    chart_type = st.sidebar.radio("Тип графика ", ("гистограмма", "диаграмма"))

    if chart_type == "гистограмма":
        # Подготовка данных для гистограммы
        bar_data = results_df.copy()
        bar_data['Время'] = pd.to_timedelta(bar_data['Время']).dt.total_seconds() / 3600

        # Создайте столбчатую Штабелированная гистограмма для Общего времени для каждого разработчика
        fig_bar = px.bar(bar_data, x='Разработчики', y='Время', title='гистограмма для каждого разработчика')
        # Показать гистограмму
        st.plotly_chart(fig_bar)

        # Подготовка данных для второй гистограммы
        cleaned_bar_data = cleaned_results_df.copy()
        cleaned_bar_data['Время'] = pd.to_timedelta(cleaned_bar_data['Время']).dt.total_seconds() / 3600

        # Создайте столбчатую Штабелированная гистограмма для Общего времени для каждого разработчика (без времени отдыха < 30 минут)
        fig_cleaned_bar = px.bar(cleaned_bar_data, x='Разработчики', y='Время', title='гистограмма для каждого разработчика (без времени отдыха < 30 минут)')
        # Показать гистограмму
        st.plotly_chart(fig_cleaned_bar)

    elif chart_type == "диаграмма":
        # Подготовьте данные для круговой диаграммы
        pie_data = results_df.copy()
        pie_data['Время'] = pd.to_timedelta(pie_data['Время']).dt.total_seconds()

        # Создайте круговую диаграмму для Общего времени для каждого разработчика
        fig_pie = px.pie(pie_data, values='Время', names='Разработчики', title='диаграмма для каждого разработчика')
        # Задайте значение для изменения расстояния между элементами в круговой окружности
        fig_pie.update_traces(hole=0.3)
        # Показать круговую диаграмму
        st.plotly_chart(fig_pie)

        # Подготовьте данные для второй круговой диаграммы
        pie_cleaned_data = cleaned_results_df.copy()
        pie_cleaned_data['Время'] = pd.to_timedelta(pie_cleaned_data['Время']).dt.total_seconds()

        # Создайте круговую диаграмму для Общего времени для каждого разработчика (без времени отдыха < 30 минут)
        fig_cleaned_pie = px.pie(pie_cleaned_data, values='Время', names='Разработчики', title='диаграмма для каждого разработчика (без времени отдыха < 30 минут)')
        # Задайте значение для изменения расстояния между элементами в круговой окружности
        fig_cleaned_pie.update_traces(hole=0.3)
        # Показать круговую диаграмму
        st.plotly_chart(fig_cleaned_pie)
              
#------------------------------------------------------------------------------------------------------------------------------------      
#------------------------------------------------------------------------------------------------------------------------------------      
#------------------------------------------------------------------------------------------------------------------------------------          

# Расчет Количество элементов          
st.sidebar.title("Расчет Количество элементов")
if st.sidebar.checkbox("элементов"):
    uploaded_files = st.sidebar.file_uploader("Количество элементов", type=['csv'], accept_multiple_files=True, key='file_uploader', help="Загрузить файлы .csv")

    # Импортируйте данные из каждого файла и добавьте их в список
    if uploaded_files is not None:
        
        all_data = []
        total_sum_all_files = 0  # Общее количество для всех файлов
        
        for file in uploaded_files:
            df = pd.read_csv(file)
            # Основные данные перед анализом
            if 'Количество' not in df.columns:
                df['Количество'] = 1
            sum_values = df['Количество'].sum()
            total_sum_all_files += sum_values  # Добавление к общему количеству
            file_name = file.name.split('.')[0]  # Имя файла без расширения
            data = {
                'Элемент': [file_name],
                'количество': [sum_values]
            }
            df_result = pd.DataFrame(data)
            all_data.append(df_result)
            
        # Объединить все данные в одну таблицу
        if len(all_data) > 0:
            combined_df = pd.concat(all_data, ignore_index=True)
            
            # Создать новый DataFrame для отображения общего количества один раз
            total_df = pd.DataFrame({'Элемент': ['Общее количество'], 'количество': [total_sum_all_files]})
            
            # Добавить общее количество в общую таблицу
            combined_df = pd.concat([combined_df, total_df], ignore_index=True)
            
            # Отображение таблиц
            st.write('Количество элементов:')
            st.write(combined_df)
            
            # Добавьте кнопку для сохранения таблицы в файле Word
            if st.button("Сохранение таблицы в файл Word"):
                # Создайте файл Word и запишите в него таблицу
                document = Document()
                
               # Добавьте изображение в верхнюю часть документа
                document.add_picture('Урфу.jpg', width=Inches(1.0), height=Inches(1.0))
                
                document.add_heading('Таблица элементов', level=1)

                # Добавить таблицу
                table = document.add_table(rows=combined_df.shape[0]+1, cols=combined_df.shape[1])
                # Добавить заголовки
                for j in range(combined_df.shape[-1]):
                    table.cell(0, j).text = combined_df.columns[j]
                # Добавить данные
                for i in range(combined_df.shape[0]):
                    for j in range(combined_df.shape[-1]):
                        table.cell(i+1, j).text = str(combined_df.iloc[i,j])

                # Сохраните файл
                document.save('таблица элементов.docx')
                st.success("Таблица была успешно сохранена в файле Word.")
               


                
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------

#Добавление новой страницы "О программе "
st.sidebar.title("О программе")
if st.sidebar.checkbox("О программе"):
  st.sidebar.write("-  Это программа Была создана командой номер 32 .")
  st.sidebar.write("-  (Главный куратор) -Машкин Олег Владимирович ")
  st.sidebar.write("- (Тимлид) - Бадри Хазем Хешам Мухаммед")       
  

#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------

#Добавление новой страницы "Renga Api"
if st.sidebar.checkbox("Renga Api"):
    st.sidebar.image("renga.png", width=80, use_column_width=False)
    st.sidebar.markdown("[Открыть](https://help.rengabim.com/api/)")
    

#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------


# #Открыть файл Renga.rnp

#Открыть файл Renga.rnp
if st.sidebar.checkbox("Открыть файл Renga.rnp"):
    uploaded_file = st.sidebar.file_uploader("Импорт файла Renga", type=["rnp"])
    if uploaded_file is not None:
        # Create the 'temp' directory if it doesn't exist
        if not os.path.exists("temp"):
            os.makedirs("temp")
        
        # Save the uploaded file in the 'temp' directory
        with open(os.path.join("temp", uploaded_file.name), "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Запускаем файл
        os.startfile(os.path.join("temp", uploaded_file.name))
        st.sidebar.success("Файл был успешно импортирован и запущен!")



#-----------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------

#Помощь
if st.sidebar.checkbox("Помочь?"):
  # Просматривайте подсказки с помощью markdown
  st.sidebar.title("Как работать:")
  st.sidebar.write("- Выберите свой файл в формате журнала.LOG")


#-----------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------------------------------------------------------

